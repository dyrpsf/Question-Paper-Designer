import customtkinter as ctk
from tkinter import messagebox, filedialog
from bs4 import BeautifulSoup, NavigableString, Tag
import re
import traceback
import time
import threading
import tempfile
import os
import io
from urllib.parse import urljoin

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from webdriver_manager.chrome import ChromeDriverManager

def _ensure_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        return Document, Pt, Inches
    except ModuleNotFoundError:
        messagebox.showerror(
            "Missing dependency",
            "The package 'python-docx' is required to export DOCX.\n\nInstall it with:\n\npip install python-docx"
        )
        return None, None, None

def _ensure_requests():
    try:
        import requests
        return requests
    except ModuleNotFoundError:
        messagebox.showerror(
            "Missing dependency",
            "The package 'requests' is required to download images for diagrams.\n\nInstall it with:\n\npip install requests"
        )
        return None

def _ensure_pillow():
    try:
        from PIL import Image
        return Image
    except ModuleNotFoundError:
        messagebox.showerror(
            "Missing dependency",
            "The package 'Pillow' is required to preview/convert images.\n\nInstall it with:\n\npip install pillow"
        )
        return None

class QuestionPaperApp(ctk.CTk):

    MAX_IMAGES_PER_QUESTION = 6
    OPTION_CAPTURE_WINDOW = 12
    OPTION_LABEL_STYLE = "a)"

    def __init__(self):
        super().__init__()

        self.title("📄 Question Paper Generator")
        self.geometry("950x740")
        self.iconbitmap("qp_ico.ico")

        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")

        self.container = ctk.CTkFrame(self)
        self.container.pack(fill="both", expand=True, padx=10, pady=(10, 0))

        self.status_frame = ctk.CTkFrame(self)
        self.status_frame.pack(fill="x", padx=10, pady=(5, 10))

        self.status_label = ctk.CTkLabel(self.status_frame, text="Ready", anchor="w")
        self.status_label.pack(side="left", padx=(0, 10))

        self.loader = ctk.CTkProgressBar(self.status_frame, mode='indeterminate', width=220)
        self.loader.pack_forget()

        self.checkbox_vars = []
        self.checkboxes = []
        self.url_entries = []
        self.all_questions = []
        self.current_disable_widget = None

        self._preview_temp_files = []

        self._show_start_screen()

        self.protocol("WM_DELETE_WINDOW", self._on_app_close)

    def _clear_container(self):
        for child in self.container.winfo_children():
            child.destroy()

    def _set_status(self, text):
        self.status_label.configure(text=text)

    def _start_loader(self, disable_widget=None):
        self.current_disable_widget = disable_widget
        if disable_widget is not None:
            try:
                disable_widget.configure(state="disabled")
            except Exception:
                pass
        self.loader.pack(side="left", padx=(10, 0))
        self.loader.start()

    def _stop_loader(self):
        self.loader.stop()
        self.loader.pack_forget()
        if self.current_disable_widget is not None:
            try:
                self.current_disable_widget.configure(state="normal")
            except Exception:
                pass
            self.current_disable_widget = None

    def _on_app_close(self):
        for p in self._preview_temp_files:
            try:
                os.remove(p)
            except Exception:
                pass
        self.destroy()

    def _show_start_screen(self):
        self._clear_container()

        frame = ctk.CTkFrame(self.container)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        title = ctk.CTkLabel(frame, text="How many websites do you want to add?", font=ctk.CTkFont(size=18, weight="bold"))
        title.pack(pady=(20, 10))

        self.count_entry = ctk.CTkEntry(frame, placeholder_text="Enter a number (e.g., 3)", width=300, justify="center")
        self.count_entry.pack(pady=(0, 15))

        btn = ctk.CTkButton(frame, text="➡️ Next", command=self._go_to_url_inputs)
        btn.pack(pady=(0, 20))

        tip = ctk.CTkLabel(frame, text="Paste multiple URLs, merge & dedupe questions, and export to one DOCX.\nDiagrams/photos near questions will be included.", justify="center")
        tip.pack()

    def _go_to_url_inputs(self):
        raw = (self.count_entry.get() or "").strip()
        if not raw.isdigit():
            messagebox.showerror("Invalid input", "Please enter a valid positive number.")
            return
        count = int(raw)
        if count <= 0 or count > 20:
            messagebox.showerror("Invalid input", "Please enter a number between 1 and 20.")
            return

        self._show_url_input_screen(count)

    def _show_url_input_screen(self, count):
        self._clear_container()
        self.url_entries.clear()

        frame = ctk.CTkFrame(self.container)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        header = ctk.CTkLabel(frame, text=f"Enter {count} website URL(s)", font=ctk.CTkFont(size=18, weight="bold"))
        header.pack(pady=(10, 5))

        url_scroll = ctk.CTkScrollableFrame(frame, width=880, height=360)
        url_scroll.pack(pady=(10, 10), fill="both", expand=True)

        for i in range(count):
            entry = ctk.CTkEntry(url_scroll, placeholder_text=f"Paste website URL #{i+1} (with questions)", width=850)
            entry.pack(pady=6)
            self.url_entries.append(entry)

        button_row = ctk.CTkFrame(frame)
        button_row.pack(pady=(10, 6))

        back_btn = ctk.CTkButton(button_row, text="⬅️ Back", width=140, command=self._show_start_screen)
        back_btn.pack(side="left", padx=10)

        self.load_all_button = ctk.CTkButton(button_row, text="🔍 Load Questions", width=180, command=self.load_questions_async_multi)
        self.load_all_button.pack(side="left", padx=10)

        tip = ctk.CTkLabel(frame, text="Notes:\n- All questions across these URLs will be merged and deduplicated.\n- Select the questions you want and export to one DOCX file.\n- Diagrams/photos located near questions will be attached.", justify="left")
        tip.pack(pady=(6, 10), anchor="w", padx=4)

    def _show_results_screen(self):
        self._clear_container()

        frame = ctk.CTkFrame(self.container)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        header = ctk.CTkLabel(frame, text="Select Questions to Export", font=ctk.CTkFont(size=18, weight="bold"))
        header.pack(pady=(10, 5))

        controls = ctk.CTkFrame(frame)
        controls.pack()

        self.select_all_btn = ctk.CTkButton(controls, text="✅ Select All", width=140, command=self.select_all)
        self.select_all_btn.pack(side="left", padx=10)

        self.deselect_all_btn = ctk.CTkButton(controls, text="❌ Deselect All", width=140, command=self.deselect_all)
        self.deselect_all_btn.pack(side="left", padx=10)

        self.found_label = ctk.CTkLabel(controls, text="")
        self.found_label.pack(side="left", padx=10)

        self.scroll_frame = ctk.CTkScrollableFrame(frame, width=880, height=460)
        self.scroll_frame.pack(pady=12, fill="both", expand=True)

        self.export_button = ctk.CTkButton(frame, text="📤 Export Selected Questions to DOCX (with diagrams)", command=self.export_to_docx)
        self.export_button.pack(pady=10)

        back_btn = ctk.CTkButton(frame, text="⬅️ Back to URLs", command=lambda: self._show_url_input_screen(len(self.url_entries)))
        back_btn.pack(pady=(0, 10))

    def load_questions_async_multi(self):
        self.checkbox_vars.clear()
        self.checkboxes.clear()

        urls = []
        for e in self.url_entries:
            u = (e.get() or "").strip()
            if u:
                urls.append(u)

        if not urls:
            messagebox.showerror("❌ No URLs", "Please enter at least one website URL.")
            return

        invalids = [u for u in urls if not u.startswith("http")]
        if invalids:
            messagebox.showerror("❌ Invalid URLs", "Please ensure all URLs start with http or https.")
            return

        self._set_status(f"Starting to load {len(urls)} site(s)...")
        self._start_loader(disable_widget=self.load_all_button)
        threading.Thread(target=self._load_questions_worker_multi, args=(urls,), daemon=True).start()

    def _load_questions_worker_multi(self, urls):
        aggregated = []
        seen_map = {}
        errors = []

        for idx, url in enumerate(urls, start=1):
            try:
                self.after(0, self._set_status, f"Loading {idx}/{len(urls)}: {url}")
                html = self._scrape_page(url)
                qdatas = self._extract_questions(html, url)
                for qd in qdatas:
                    key = (qd.get("key") or qd["text"]).lower()
                    if key in seen_map:
                        i = seen_map[key]
                        existing = set(aggregated[i]["images"])
                        for u in qd["images"]:
                            if u not in existing and len(aggregated[i]["images"]) < self.MAX_IMAGES_PER_QUESTION:
                                aggregated[i]["images"].append(u)
                                existing.add(u)
                    else:
                        seen_map[key] = len(aggregated)
                        qd["images"] = qd["images"][: self.MAX_IMAGES_PER_QUESTION]
                        aggregated.append(qd)
            except Exception as e:
                err = f"{e.__class__.__name__}: {e}\n{traceback.format_exc()}"
                errors.append((url, err))

        self.after(0, self._on_questions_loaded_multi, aggregated, errors)

    def _on_questions_loaded_multi(self, questions, errors):
        self._stop_loader()

        if errors:
            msg = "Some URLs could not be loaded:\n\n"
            for url, err in errors[:3]:
                msg += f"- {url}\n  {err.splitlines()[0]}\n"
            if len(errors) > 3:
                msg += f"... and {len(errors)-3} more.\n"
            messagebox.showwarning("⚠️ Partial Errors", msg)

        if not questions:
            messagebox.showinfo("❗ No Questions", "No valid questions found across the provided URLs.")
            self._set_status("No questions found.")
            return

        self.all_questions = questions
        total_imgs = sum(len(q["images"]) for q in self.all_questions)
        self._set_status(f"Found {len(questions)} question(s), {total_imgs} diagram(s). Rendering list...")

        self._show_results_screen()

        for qd in self.all_questions:
            self._add_question_row(self.scroll_frame, qd)

        if hasattr(self, "found_label"):
            self.found_label.configure(text=f"Found {len(self.all_questions)} question(s) • {total_imgs} diagram(s)")

        self._set_status("Ready")

    def _add_question_row(self, parent, qdata):
        var = ctk.BooleanVar(value=False)

        row = ctk.CTkFrame(parent)
        row.pack(fill="x", pady=(6, 0))

        cb = ctk.CTkCheckBox(row, text="", variable=var, width=24)
        cb.pack(side="left", padx=(0, 8))
        cb._qdata = qdata

        label_text = qdata["text"]
        lbl = ctk.CTkLabel(row, text=label_text, wraplength=710, justify="left", anchor="w")
        lbl.pack(side="left", fill="x", expand=True)

        if qdata.get("images"):
            prev_btn = ctk.CTkButton(row, text=f"Preview ({len(qdata['images'])})", width=120,
                                     command=lambda r=row, q=qdata: self._toggle_preview(r, q))
            prev_btn.pack(side="right", padx=6)

        preview_frame = ctk.CTkFrame(parent)
        preview_frame.pack(fill="x", padx=26, pady=(4, 6))
        preview_frame.pack_forget()

        row._preview_frame = preview_frame
        row._preview_loaded = False
        row._preview_visible = False
        row._thumb_refs = []
        row._thumb_paths = []

        self.checkbox_vars.append(var)
        self.checkboxes.append(cb)

        sep = ctk.CTkFrame(parent, height=1)
        sep.pack(fill="x", padx=26, pady=(6, 0))

    def _toggle_preview(self, row, qdata):
        if row._preview_visible:
            row._preview_frame.pack_forget()
            row._preview_visible = False
            return

        if not row._preview_loaded:
            self._load_preview_async(row, qdata)
        else:
            row._preview_frame.pack(fill="x", padx=26, pady=(4, 6))
            row._preview_visible = True

    def _load_preview_async(self, row, qdata):
        requests = _ensure_requests()
        ImageLib = _ensure_pillow()
        if requests is None or ImageLib is None:
            return

        for child in row._preview_frame.winfo_children():
            child.destroy()
        loading = ctk.CTkLabel(row._preview_frame, text="Loading diagrams...")
        loading.pack(pady=8)
        row._preview_frame.pack(fill="x", padx=26, pady=(4, 6))
        row._preview_visible = True

        def worker():
            thumbs = []
            for url in qdata.get("images", [])[: self.MAX_IMAGES_PER_QUESTION]:
                path, was_webp, conv_failed = self._download_image(url, requests, ImageLib)
                if not path:
                    continue
                try:
                    img = ImageLib.open(path)
                    img = img.convert("RGB") if img.mode not in ("RGB", "L") else img
                    max_w, max_h = 220, 160
                    img.thumbnail((max_w, max_h))
                    w, h = img.size
                    ctk_img = ctk.CTkImage(light_image=img, size=(w, h))
                    thumbs.append((ctk_img, path))
                except Exception:
                    try:
                        os.remove(path)
                    except Exception:
                        pass
                    continue

            def on_ui():
                for child in row._preview_frame.winfo_children():
                    child.destroy()

                if not thumbs:
                    ctk.CTkLabel(row._preview_frame, text="No diagrams to preview.").pack(pady=8)
                else:
                    grid = ctk.CTkFrame(row._preview_frame)
                    grid.pack(fill="x", padx=4, pady=4)
                    cols = 3
                    for i, (ctk_img, path) in enumerate(thumbs):
                        r = i // cols
                        c = i % cols
                        lbl = ctk.CTkLabel(grid, image=ctk_img, text="")
                        lbl.grid(row=r, column=c, padx=6, pady=6, sticky="w")
                        row._thumb_refs.append(ctk_img)
                        row._thumb_paths.append(path)
                        self._preview_temp_files.append(path)

                row._preview_loaded = True
                row._preview_visible = True

            self.after(0, on_ui)

        threading.Thread(target=worker, daemon=True).start()

    def _scrape_page(self, url):
        options = Options()
        try:
            options.add_argument("--headless=new")
        except Exception:
            options.add_argument("--headless")
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--window-size=1280,1800")
        options.add_argument("--log-level=3")
        options.add_argument("--disable-extensions")
        options.page_load_strategy = "eager"

        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

        try:
            try:
                driver.execute_cdp_cmd("Network.enable", {})
                driver.execute_cdp_cmd("Network.setBlockedURLs", {
                    "urls": ["*.mp4", "*.webm", "*.avi", "*.mov", "*.m3u8"]
                })
            except Exception:
                pass

            driver.set_page_load_timeout(50)
            driver.get(url)

            WebDriverWait(driver, 25).until(
                lambda d: len(d.find_elements(By.CSS_SELECTOR, "p, li")) >= 3 or len(d.page_source) > 20000
            )

            self._expand_all(driver)
            self._progressive_scroll(driver)
            time.sleep(0.6)

            return driver.page_source
        finally:
            driver.quit()

    def _progressive_scroll(self, driver, steps=12, pause=0.35):
        try:
            last_height = driver.execute_script("return document.body.scrollHeight || document.documentElement.scrollHeight;")
            for i in range(steps):
                y = int((i / float(max(steps - 1, 1))) * last_height)
                driver.execute_script("window.scrollTo(0, arguments[0]);", y)
                time.sleep(pause)
                new_height = driver.execute_script("return document.body.scrollHeight || document.documentElement.scrollHeight;")
                if new_height > last_height:
                    last_height = new_height
            driver.execute_script("window.scrollBy(0, -200);")
        except Exception:
            pass

    def _expand_all(self, driver):
        try:
            elements = driver.find_elements(By.CLASS_NAME, "accordion-item")
            for el in elements:
                try:
                    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
                except Exception:
                    pass

                toggle = None
                toggle_selectors = [
                    "[data-bs-toggle='collapse']",
                    ".accordion-header button",
                    "button[aria-expanded]",
                    "button",
                    "[role='button']",
                    ".accordion-header",
                    "summary",
                ]
                for sel in toggle_selectors:
                    try:
                        toggle = el.find_element(By.CSS_SELECTOR, sel)
                        break
                    except Exception:
                        continue

                if toggle is None:
                    toggle = el

                try:
                    expanded = (toggle.get_attribute("aria-expanded") or "").lower() == "true"
                except Exception:
                    expanded = False

                if not expanded:
                    try:
                        driver.execute_script("arguments[0].click();", toggle)
                    except Exception:
                        try:
                            toggle.click()
                        except Exception:
                            pass
                    time.sleep(0.2)
        except Exception as e:
            print("Accordion expansion issue:", e)

        try:
            driver.execute_script("document.querySelectorAll('details').forEach(d => d.open = true);")
        except Exception:
            pass

        try:
            for hdr in driver.find_elements(By.CSS_SELECTOR, ".mat-expansion-panel-header, mat-expansion-panel .mat-expansion-panel-header"):
                try:
                    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", hdr)
                    driver.execute_script("arguments[0].click();", hdr)
                    time.sleep(0.1)
                except Exception:
                    pass
        except Exception:
            pass

        try:
            for el in driver.find_elements(By.CSS_SELECTOR, ".v-expansion-panel-title, .v-expansion-panel__header"):
                try:
                    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
                    driver.execute_script("arguments[0].click();", el)
                    time.sleep(0.1)
                except Exception:
                    pass
        except Exception:
            pass

        try:
            candidates = driver.find_elements(By.XPATH, "//button|//a|//div")
            for c in candidates:
                try:
                    txt = (c.text or "").strip().lower()
                    if any(k in txt for k in ["show more", "read more", "expand", "view all", "show all"]):
                        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", c)
                        driver.execute_script("arguments[0].click();", c)
                        time.sleep(0.15)
                except Exception:
                    continue
        except Exception:
            pass

    def _looks_like_question_text(self, text: str) -> bool:
        t = (text or "").strip()
        if not t:
            return False
        tl = t.lower()
        if t.endswith("?"):
            return True
        if re.match(r"^\s*(?:q(?:uestion)?\s*[:.]?\s*|\(?\d{1,3}\)?\s*[.)]\s*)", t, re.I):
            return True
        if re.match(r"^(what|which|when|where|why|how|explain|define|calculate|find|determine|state|solve|write|prove|show|discuss|differentiate|evaluate|program|integrate|mention|draw|give|example|compare|convert|display|predict)\b", t, re.I):
            return True
        if "prove that" in tl or "show that" in tl:
            return True
        return False

    def _looks_like_option_line(self, text: str) -> bool:
        if not text:
            return False
        s = text.strip()
        if not s:
            return False

        if (re.match(r"^\s*q(?:uestion|n)?\b", s, re.I) or re.match(r"^\s*\(?\d{1,3}\)?\s*[.)]", s)) and s.endswith("?"):
            return False

        if re.match(r"^\s*[\(\[]?\s*[A-Za-z]\s*[\)\].:-]\s+\S", s):
            return True
        if re.match(r"^\s*[\(\[]?\s*[ivxlcdm]{1,7}\s*[\)\].:-]\s+\S", s, re.I):
            return True
        if re.match(r"^\s*[•\-\u2013\u2014]\s+\S", s):
            return True

        return False

    def _append_option_line(self, qdata, line: str):
        line = (line or "").strip()
        if not line:
            return
        if not qdata.get("text"):
            qdata["text"] = line
            return
        last_line = qdata["text"].splitlines()[-1].strip() if qdata["text"] else ""
        if line.lower() == last_line.lower():
            return
        qdata["text"] += "\n" + line

    def _strip_leading_enumeration(self, text: str) -> str:
        t = text or ""
        t = re.sub(r"^\s*q(?:uestion)?\s*[:.\-#]?\s*\d+\s*[:.)-]*\s*", "", t, flags=re.I)
        t = re.sub(r"^\s*que?\.\s*no\.\s*\d+\s*[:.)-]*\s*", "", t, flags=re.I)
        t = re.sub(r"^\s*q[n]?\s*[-:.#]?\s*\d+\s*[:.)-]*\s*", "", t, flags=re.I)
        t = re.sub(r"^\s*\(?\d{1,3}\)?\s*[.)-:]\s*", "", t)
        t = re.sub(r"^\s*\(?[a-zA-Z]\)?\s*[.)-:]\s*", "", t)
        t = re.sub(r"^\s*\(?[ivxlcdmIVXLCDM]{1,7}\)?\s*[.)-:]\s*", "", t)
        t = re.sub(r"^\s*[•\-\u2013\u2014]\s+", "", t)
        return t.strip()

    def _extract_questions(self, html, base_url):
        soup = BeautifulSoup(html, "html.parser")
        texts = []

        for el in soup.find_all(text=True):
            parent = el.parent
            if isinstance(el, NavigableString):
                s = str(el).strip()
                if s and parent.name not in ["script", "style"]:
                    texts.append((s, parent))

        questions = []
        current_q = None
        prev_parent = None

        for text, parent in texts:
            if self._looks_like_question_text(text):
                if current_q:
                    questions.append(current_q)
                cleaned = self._strip_leading_enumeration(text)
                current_q = {"text": cleaned, "images": [], "key": cleaned.lower()}
                prev_parent = parent
                continue

            if current_q:
                parent_chain_match = False
                p = parent
                while p and p is not prev_parent:
                    p = p.parent
                if p is prev_parent:
                    parent_chain_match = True

                if parent_chain_match or abs(parent.sourceline - prev_parent.sourceline if hasattr(parent, 'sourceline') and hasattr(prev_parent, 'sourceline') else 0) <= self.OPTION_CAPTURE_WINDOW:
                    if self._looks_like_option_line(text):
                        self._append_option_line(current_q, text)
                        continue

                if parent_chain_match:
                    imgs = parent.find_all("img")
                    for im in imgs:
                        src = im.get("src") or im.get("data-src") or ""
                        if src:
                            questions[-1]["images"].append(urljoin(base_url, src))

            prev_parent = parent

        if current_q:
            questions.append(current_q)

        return questions

    def _download_image(self, url, requests, ImageLib):
        try:
            r = requests.get(url, timeout=10)
            if r.status_code != 200:
                return None, False, False

            ext = os.path.splitext(url)[1].lower()
            is_webp = (ext == ".webp")

            fd, path = tempfile.mkstemp(suffix=ext if not is_webp else ".png")
            os.close(fd)

            with open(path, "wb") as f:
                f.write(r.content)

            if is_webp:
                try:
                    img = ImageLib.open(path).convert("RGB")
                    new_path = path.replace(".png", "_conv.jpg")
                    img.save(new_path, "JPEG")
                    os.remove(path)
                    return new_path, True, False
                except Exception:
                    return path, True, True

            return path, False, False
        except Exception:
            return None, False, False

    def select_all(self):
        for var in self.checkbox_vars:
            var.set(True)

    def deselect_all(self):
        for var in self.checkbox_vars:
            var.set(False)

    def export_to_docx(self):
        Document, Pt, Inches = _ensure_docx()
        if Document is None:
            return

        selected = []
        for var, cb in zip(self.checkbox_vars, self.checkboxes):
            if var.get():
                selected.append(cb._qdata)

        if not selected:
            messagebox.showinfo("No Selection", "Please select at least one question.")
            return

        file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file:
            return

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        for i, qd in enumerate(selected, start=1):
            doc.add_paragraph(f"Q{i}. {qd['text']}")
            for url in qd["images"]:
                try:
                    requests = _ensure_requests()
                    if requests is None:
                        continue
                    r = requests.get(url, timeout=10)
                    if r.status_code != 200:
                        continue
                    fd, path = tempfile.mkstemp(suffix=os.path.splitext(url)[1])
                    os.close(fd)
                    with open(path, "wb") as f:
                        f.write(r.content)
                    doc.add_picture(path, width=Inches(4.5))
                    os.remove(path)
                except Exception:
                    continue
            doc.add_paragraph("")

        doc.save(file)
        messagebox.showinfo("Success", f"Exported {len(selected)} questions to {file}")

if __name__ == "__main__":
    app = QuestionPaperApp()
    app.mainloop()
